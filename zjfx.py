# -*- coding: utf-8 -*-
import streamlit as st
import pandas as pd
import numpy as np
import seaborn as sns
import matplotlib
import matplotlib.pyplot as plt
import statsmodels.api as sm
from scipy import stats
from itertools import combinations
from math import sqrt
from io import BytesIO
import base64
import zipfile
import datetime as dt
import os

# ML / utils
from sklearn.preprocessing import StandardScaler, PolynomialFeatures
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import roc_auc_score, average_precision_score, roc_curve, precision_recall_curve
from sklearn.neighbors import NearestNeighbors

matplotlib.rcParams['axes.unicode_minus'] = False

st.set_page_config(page_title="质检与满意度分析", layout="wide")
st.title("质检-满意度分析")

# ====================== 图表/表格注册器（用于导出） ======================
fig_registry = {}   # {name: matplotlib.figure.Figure}
table_registry = {} # {name: pandas.DataFrame}
text_registry = {}  # {name: str}

def register_fig(name, fig):
    fig_registry[name] = fig

def register_table(name, df):
    table_registry[name] = df.copy()

def register_text(name, txt):
    text_registry[name] = str(txt)

# ============ 小工具函数 ============
def safe_to_int(x):
    if pd.isna(x): return 0
    s = str(x).strip().lower()
    if s in ["1", "1.0", "true", "t", "y", "yes"]: return 1
    if s in ["0", "0.0", "false", "f", "n", "no"]: return 0
    return 0

def plot_spc(series, title="SPC Chart"):
    x = np.arange(len(series))
    mean = np.mean(series)
    std = np.std(series, ddof=1) if len(series) > 1 else 0
    ucl = mean + 3*std
    lcl = mean - 3*std
    fig, ax = plt.subplots(figsize=(9,4.5), dpi=140)
    ax.plot(x, series, marker="o", linewidth=1.8, label="Rate")
    ax.axhline(mean, linestyle="--", color="gray", label="Center")
    if std>0:
        ax.axhline(ucl, linestyle="--", color="red", label="UCL (+3σ)")
        ax.axhline(lcl, linestyle="--", color="red", label="LCL (-3σ)")
    ax.set_title(title)
    ax.set_xlabel("Index")
    ax.set_ylabel("Rate")
    ax.legend()
    ax.grid(alpha=0.25, linestyle="--", linewidth=0.5)
    return fig

def benjamini_hochberg(pvals, alpha=0.05):
    p = np.array(pvals)
    n = len(p)
    order = np.argsort(p)
    ranked = p[order]
    crit = alpha * np.arange(1, n+1) / n
    passed = ranked <= crit
    if not passed.any():
        return np.zeros(n, dtype=bool)
    k = np.max(np.where(passed))
    thresh = ranked[k]
    return p <= thresh

def bootstrap_diff_rate(a, b, n_boot=1000, seed=42):
    rng = np.random.default_rng(seed)
    diffs = []
    a = np.array(a); b = np.array(b)
    for _ in range(n_boot):
        aa = rng.choice(a, size=len(a), replace=True)
        bb = rng.choice(b, size=len(b), replace=True)
        diffs.append(aa.mean() - bb.mean())
    lo, hi = np.percentile(diffs, [2.5, 97.5])
    return np.mean(diffs), lo, hi

# ====================== 上传 ======================
uploaded_files = st.file_uploader(
    "请上传多个质检数据文件（支持 Excel / CSV，可多选）",
    type=["xlsx", "csv"],
    accept_multiple_files=True
)

if uploaded_files:
    # 读取
    all_dfs = []
    for uploaded_file in uploaded_files:
        try:
            if uploaded_file.name.lower().endswith(".csv"):
                df_tmp = pd.read_csv(uploaded_file)
            else:
                df_tmp = pd.read_excel(uploaded_file)
            df_tmp["source_file"] = uploaded_file.name
            all_dfs.append(df_tmp)
        except Exception as e:
            st.error(f"❌ 文件 {uploaded_file.name} 读取失败: {e}")
    df = pd.concat(all_dfs, ignore_index=True)
    st.success(f"✅ 成功加载 {len(uploaded_files)} 个文件，共 {len(df)} 条记录")

    # 清洗
    st.subheader("数据清洗与逻辑处理")
    required_cols = ["score", "solution", "service_attitude", "response_speed", "case_classification"]
    for c in required_cols:
        if c not in df.columns:
            st.error(f"❌ 缺少字段：{c}")
            st.stop()
    df = df.replace("-", np.nan)
    df = df[df["score"].notna()].copy()
    df["response_speed"] = df["response_speed"].fillna(1)

    pass_cols = ["solution", "service_attitude", "response_speed", "case_classification"]
    for col in pass_cols:
        df[col] = df[col].apply(safe_to_int)
    df["satisfied"] = np.where(df["score"] >= 4, 1, 0)
    df["overall_pass"] = np.where((df[pass_cols].sum(axis=1) == 4), 1, 0)

    st.markdown("""
    **清洗逻辑：** 剔除未打分；'-' 视为空；`response_speed` 空按通过；四项全 1 为整体通过；打分≥4 为满意。
    """)

    # 概览
    st.subheader("各项通过率与满意率")
    summary = df[pass_cols + ["overall_pass", "satisfied"]].mean().to_frame("rate")
    summary["rate"] = summary["rate"].apply(lambda x: round(x*100,2))
    st.dataframe(summary.T.style.format("{:.2f}%"))
    register_table("summary_rates", summary)

    # 相关性
    st.subheader("相关性分析（Pearson）")
    fig_corr, ax_corr = plt.subplots(figsize=(6.5, 4.5), dpi=140)
    sns.heatmap(df[pass_cols + ["satisfied"]].corr(), annot=True, cmap="YlGnBu", fmt=".2f", ax=ax_corr)
    ax_corr.set_title("Correlation between QC Items and Satisfaction", fontsize=11)
    st.pyplot(fig_corr)
    register_fig("corr_matrix", fig_corr)

    # t检验
    st.subheader("显著性差异分析（t 检验：通过组 vs 未通过组满意率）")
    t_rows = []
    for col in pass_cols:
        p1 = df[df[col]==1]["satisfied"]; p0 = df[df[col]==0]["satisfied"]
        if len(p1)>2 and len(p0)>2:
            t, p = stats.ttest_ind(p1, p0, equal_var=False)
            diff = p1.mean()-p0.mean()
            t_rows.append([col, round(p1.mean(),3), round(p0.mean(),3), round(diff,3), round(p,4)])
    result_df = pd.DataFrame(t_rows, columns=["指标项","通过组满意率","未通过组满意率","差异","p值"])
    st.dataframe(result_df)
    register_table("t_test", result_df)

    # Logistic 单项
    st.subheader("Logistic 回归分析（单项）")
    X = sm.add_constant(df[pass_cols]); y = df["satisfied"]
    logit_model = sm.Logit(y, X).fit(disp=False)
    coef_df = pd.DataFrame({
        "指标项": logit_model.params.index[1:],
        "回归系数": logit_model.params.values[1:],
        "p值": logit_model.pvalues.values[1:]
    }).sort_values("回归系数", ascending=False)
    st.dataframe(coef_df.style.background_gradient(cmap="RdYlGn", axis=0))
    register_table("logit_single", coef_df)

    fig_bar, ax_bar = plt.subplots(figsize=(7.5, 4.5), dpi=140)
    sns.barplot(x="回归系数", y="指标项", data=coef_df, ax=ax_bar)
    ax_bar.axvline(0, color="gray", linestyle="--")
    ax_bar.set_xlabel("Regression Coefficient", fontsize=10)
    ax_bar.set_ylabel("QC Item", fontsize=10)
    ax_bar.set_title("Impact of QC Items on Satisfaction", fontsize=11)
    st.pyplot(fig_bar)
    register_fig("logit_coef_bar", fig_bar)

    # 两两交互
    st.subheader("两两组合对满意度的影响（交互项分析）")
    comb_results, interaction_cols = [], []
    for i in range(len(pass_cols)):
        for j in range(i+1, len(pass_cols)):
            c1, c2 = pass_cols[i], pass_cols[j]
            name = f"{c1} × {c2}"
            df[name] = df[c1]*df[c2]
            interaction_cols.append(name)
            grp = df.groupby(name)["satisfied"].agg(["mean","count"]).reset_index()
            if len(grp)==2:
                diff = grp.loc[grp[name]==1,"mean"].values[0]-grp.loc[grp[name]==0,"mean"].values[0]
                t, p = stats.ttest_ind(df[df[name]==1]["satisfied"], df[df[name]==0]["satisfied"], equal_var=False)
                comb_results.append([name, round(grp.loc[grp[name]==1,"mean"].values[0],3),
                                     round(grp.loc[grp[name]==0,"mean"].values[0],3),
                                     round(diff,3), round(p,4)])
    combo_df = pd.DataFrame(comb_results, columns=["组合","交互通过组满意率","未交互组满意率","差异","p值"])
    st.dataframe(combo_df)
    register_table("interaction_ttest", combo_df)

    st.subheader("交互项 Logistic 回归")
    X_interact = sm.add_constant(df[pass_cols+interaction_cols])
    logit_interact = sm.Logit(y, X_interact).fit(disp=False)
    coef_inter_df = pd.DataFrame({
        "变量": logit_interact.params.index[1:],
        "回归系数": logit_interact.params.values[1:],
        "p值": logit_interact.pvalues.values[1:]
    }).sort_values("回归系数", ascending=False)
    st.dataframe(coef_inter_df.style.background_gradient(cmap="RdYlGn", axis=0))
    register_table("logit_interactions", coef_inter_df)

    # ========= 进阶模块（开关） =========
    st.markdown("---")
    st.header("进阶分析模块（按需开启）")

    # 1) 数据质量 / 缺失模式
    if st.checkbox("① 数据质量与缺失模式报告"):
        st.subheader("字段缺失率（%）")
        miss = df.isna().mean().sort_values(ascending=False)*100
        st.dataframe(miss.to_frame("missing_%").style.format("{:.2f}"))
        register_table("missing_report", miss.to_frame("missing_%"))
        if "satisfied" in df.columns:
            fig = plot_spc(df["satisfied"].rolling(20, min_periods=5).mean().dropna(), "SPC: Satisfaction Rolling Mean")
            st.pyplot(fig)
            register_fig("spc_satisfied", fig)

    # 2) 多重共线性 VIF
    if st.checkbox("② 多重共线性诊断（VIF）"):
        st.subheader("VIF 检验（>5/10 需警惕）")
        from statsmodels.stats.outliers_influence import variance_inflation_factor
        Xv = sm.add_constant(df[pass_cols])
        vifs = []
        for k in range(1, Xv.shape[1]):  # skip const
            vifs.append([Xv.columns[k], variance_inflation_factor(Xv.values, k)])
        vif_df = pd.DataFrame(vifs, columns=["变量","VIF"]).sort_values("VIF", ascending=False)
        st.dataframe(vif_df.style.format("{:.2f}"))
        register_table("vif", vif_df)

    # 3) 非线性（分段/单调）
    if st.checkbox("③ 非线性与单调性检验（分段）"):
        st.subheader("通过项计数 vs 满意率（分段）")
        df["pass_sum"] = df[pass_cols].sum(axis=1)
        g = df.groupby("pass_sum")["satisfied"].mean().reset_index()
        fig, ax = plt.subplots(figsize=(6.5,4), dpi=140)
        ax.plot(g["pass_sum"], g["satisfied"], marker="o")
        ax.set_title("Satisfaction by Number of Passed QC Items")
        ax.set_xlabel("#Passed Items"); ax.set_ylabel("Satisfaction Rate")
        st.pyplot(fig)
        register_fig("nonlinear_passsum", fig)
        register_table("passsum_curve", g)

    # 4) 正则化 + 交互搜索（L1/L2）
    if st.checkbox("④ 正则化逻辑回归（含大规模交互搜索）"):
        st.subheader("L1/L2 Logistic with Interactions")
        poly = PolynomialFeatures(degree=2, include_bias=False, interaction_only=False)
        Xp = poly.fit_transform(df[pass_cols])
        feat_names = poly.get_feature_names_out(pass_cols)
        scaler = StandardScaler()
        Xs = scaler.fit_transform(Xp)
        # L1
        lr_l1 = LogisticRegression(penalty="l1", solver="liblinear", max_iter=2000)
        lr_l1.fit(Xs, y)
        coef_l1 = pd.Series(lr_l1.coef_[0], index=feat_names).sort_values(ascending=False)
        nz = coef_l1[coef_l1!=0]
        st.write("L1 non-zero features:"); st.dataframe(nz.to_frame("coef"))
        register_table("l1_nonzero", nz.to_frame("coef"))
        # L2
        lr_l2 = LogisticRegression(penalty="l2", solver="liblinear", max_iter=2000)
        lr_l2.fit(Xs, y)
        auc = roc_auc_score(y, lr_l2.predict_proba(Xs)[:,1])
        ap = average_precision_score(y, lr_l2.predict_proba(Xs)[:,1])
        st.write(f"AUC={auc:.3f}, AP={ap:.3f}")
        register_text("l2_metrics", f"AUC={auc:.3f}, AP={ap:.3f}")

    # 5) 自举置信区间 + FDR
    if st.checkbox("⑤ 稳健性：Bootstrap CI + FDR 多重检验"):
        st.subheader("t检验差异的 Bootstrap 置信区间 & FDR 校正")
        rows = []
        for col in pass_cols:
            a = df[df[col]==1]["satisfied"]; b = df[df[col]==0]["satisfied"]
            if len(a)>20 and len(b)>20:
                diff, lo, hi = bootstrap_diff_rate(a,b, n_boot=1000)
                t, p = stats.ttest_ind(a,b, equal_var=False)
                rows.append([col, diff, lo, hi, p])
        robust_df = pd.DataFrame(rows, columns=["指标项","差异","CI_lo","CI_hi","p"])
        if not robust_df.empty:
            mask = benjamini_hochberg(robust_df["p"].values, alpha=0.05)
            robust_df["FDR_significant"] = mask
        st.dataframe(robust_df)
        register_table("bootstrap_fdr", robust_df)

    # 6) 误差域四象限
    if st.checkbox("⑥ 误差域四象限：标准与感知偏离定位"):
        st.subheader("四象限分布（聚焦 FP：通过但不满意；FN：未通过但满意）")
        quad = pd.crosstab(df["overall_pass"], df["satisfied"], normalize=True).round(3)
        st.dataframe(quad)
        register_table("quadrant", quad)

    # 7) 异质性：分层模型
    if st.checkbox("⑦ 异质性：分业务线/渠道/问题类型建模"):
        st.subheader("分层 Logistic（示例：business_line / ticket_channel / case_classification）")
        layered_out = []
        for dim in ["business_line", "ticket_channel", "case_classification"]:
            if dim in df.columns:
                for v in df[dim].dropna().unique().tolist()[:10]:
                    sub = df[df[dim]==v]
                    if len(sub)>50 and sub["satisfied"].nunique()>1:
                        Xs_ = sm.add_constant(sub[pass_cols])
                        try:
                            m = sm.Logit(sub["satisfied"], Xs_).fit(disp=False)
                            part = pd.DataFrame({"维度": dim, "取值": v,
                                                 "项": m.params.index[1:], "系数": m.params.values[1:], "p": m.pvalues.values[1:]})
                            layered_out.append(part)
                        except Exception:
                            pass
        if layered_out:
            layered_df = pd.concat(layered_out, ignore_index=True)
            st.dataframe(layered_df.sort_values(["维度","取值","系数"], ascending=[True,True,False]))
            register_table("layered_logit", layered_df)

    # 8) 因果：PSM
    if st.checkbox("⑧ 因果推断：倾向得分匹配（PSM）估计各项对满意度的因果提升"):
        st.subheader("PSM（每个QC项单独做一次处理与对照匹配）")
        ps_rows = []
        for col in pass_cols:
            covars = [c for c in pass_cols if c!=col]
            if len(covars)<1: continue
            lr = LogisticRegression(solver="liblinear")
            lr.fit(df[covars], df[col])
            ps = lr.predict_proba(df[covars])[:,1]
            data = df[[col, "satisfied"]].copy()
            data["ps"] = ps

            treat = data[data[col]==1].copy()
            ctrl = data[data[col]==0].copy()
            if len(treat)<10 or len(ctrl)<10: continue
            nn = NearestNeighbors(n_neighbors=1, metric="euclidean")
            nn.fit(ctrl[["ps"]].values)
            dist, idx = nn.kneighbors(treat[["ps"]].values)
            matched_ctrl = ctrl.iloc[idx.flatten()]
            att = treat["satisfied"].values.mean() - matched_ctrl["satisfied"].values.mean()
            ps_rows.append([col, att, len(treat), len(ctrl)])
        ps_df = pd.DataFrame(ps_rows, columns=["质检项","ATT(匹配后满意度提升)","treat_n","ctrl_n"])
        st.dataframe(ps_df.sort_values("ATT(匹配后满意度提升)", ascending=False))
        register_table("psm_att", ps_df)

    # 9) 过程质量监控
    if st.checkbox("⑨ 过程质量监控：SPC/CUSUM/EWMA"):
        st.subheader("SPC（滚动满意率）")
        fig = plot_spc(df["satisfied"].rolling(30, min_periods=10).mean().dropna(), "SPC: Rolling Satisfaction Rate")
        st.pyplot(fig)
        register_fig("spc_rolling", fig)

        st.subheader("CUSUM（上偏/下偏累计）")
        series = df["satisfied"].rolling(30, min_periods=10).mean().dropna()
        if len(series)>5:
            target = series.mean()
            pos = np.maximum(0, (series-target))
            neg = np.maximum(0, (target-series))
            pos_cusum = pos.cumsum(); neg_cusum = neg.cumsum()
            fig2, ax2 = plt.subplots(figsize=(9,4.5), dpi=140)
            ax2.plot(pos_cusum, label="Positive CUSUM"); ax2.plot(neg_cusum, label="Negative CUSUM")
            ax2.set_title("CUSUM of Satisfaction Deviation"); ax2.legend(); ax2.grid(alpha=0.25, linestyle="--")
            st.pyplot(fig2)
            register_fig("cusum", fig2)

    # 10) 目标建模：预测 + SHAP
    if st.checkbox("⑩ 目标建模：满意度预测 + 可解释性（SHAP）"):
        st.subheader("预测满意度（仅QC特征）")
        Xbin = df[pass_cols+interaction_cols].copy()
        scaler = StandardScaler()
        Xs = scaler.fit_transform(Xbin)
        clf = LogisticRegression(max_iter=2000, solver="liblinear")
        clf.fit(Xs, y)
        prob = clf.predict_proba(Xs)[:,1]
        auc = roc_auc_score(y, prob); ap = average_precision_score(y, prob)
        st.write(f"AUC={auc:.3f}, AP={ap:.3f}")
        register_text("clf_metrics", f"AUC={auc:.3f}, AP={ap:.3f}")

        fpr, tpr, _ = roc_curve(y, prob)
        fig, ax = plt.subplots(figsize=(6,4), dpi=140)
        ax.plot(fpr, tpr); ax.plot([0,1],[0,1], linestyle="--", color="gray")
        ax.set_title("ROC Curve"); ax.set_xlabel("FPR"); ax.set_ylabel("TPR")
        st.pyplot(fig); register_fig("roc", fig)

        pr, rc, _ = precision_recall_curve(y, prob)
        fig2, ax2 = plt.subplots(figsize=(6,4), dpi=140)
        ax2.plot(rc, pr); ax2.set_title("Precision-Recall Curve"); ax2.set_xlabel("Recall"); ax2.set_ylabel("Precision")
        st.pyplot(fig2); register_fig("pr_curve", fig2)

        try:
            import shap
            explainer = shap.LinearExplainer(clf, Xs)
            shap_values = explainer.shap_values(Xs)
            top_idx = np.argsort(np.abs(shap_values).mean(axis=0))[::-1][:20]
            top_features = pd.Series(np.abs(shap_values).mean(axis=0)[top_idx],
                                     index=pd.Index(np.array(pass_cols+interaction_cols)[top_idx], name="Feature"))
            st.dataframe(top_features.to_frame("mean|SHAP|"))
            register_table("shap_top", top_features.to_frame("mean|SHAP|"))
        except Exception as e:
            st.info(f"SHAP 不可用：{e}")

    # 11) 生存分析（可选）
    if st.checkbox("⑪ 生存分析（处理耗时对满意度的影响）"):
        st.subheader("若有处理时长字段（如：解决耗时hour），可做KM曲线")
        if "handle_hours" in df.columns and "resolved" in df.columns:
            try:
                from lifelines import KaplanMeierFitter
                km = KaplanMeierFitter()
                T = df["handle_hours"]; E = df["resolved"]
                fig, ax = plt.subplots(figsize=(7,4), dpi=140)
                km.fit(T, event_observed=E, label="All"); km.plot(ax=ax)
                ax.set_title("KM Curve of Resolution Time"); ax.set_xlabel("Hours"); ax.set_ylabel("Survival")
                st.pyplot(fig); register_fig("km_curve", fig)
            except Exception as e:
                st.info(f"需要 lifelines 包：{e}")
        else:
            st.info("未检测到 'handle_hours' 与 'resolved' 字段，跳过。")

    # 12) 策略优化建议（基于以上分析即时生成）
    st.subheader("📌 质检标准与抽检策略优化（自动建议）")
    try:
        sig = coef_df[coef_df["p值"]<0.05].copy()
        if not sig.empty:
            sig["权重候选"] = (sig["回归系数"].abs()/sig["回归系数"].abs().sum()).round(3)
            st.markdown("**1) 指标权重建议（归一化系数）：**")
            st.dataframe(sig[["指标项","回归系数","p值","权重候选"]])
            register_table("weight_candidate", sig[["指标项","回归系数","p值","权重候选"]])
        else:
            st.info("暂无显著单项，建议扩大样本。")

        top_inter = coef_inter_df[(coef_inter_df["p值"]<0.05) & (coef_inter_df["变量"].str.contains("×"))].head(3)
        st.markdown("**2) 组合联检（交互显著 Top3）：**")
        if not top_inter.empty:
            st.dataframe(top_inter); register_table("top_interactions", top_inter)
        else:
            st.info("暂无显著交互。")

        st.markdown("""
        **3) 抽检策略：**  
        - 对 **FP 区域（overall_pass=1，satisfied=0）** 提高复核概率（如：+30%），重点审查 `服务态度` 与 `解决方案`；  
        - 对 **FN 区域（overall_pass=0，satisfied=1）** 复盘标准是否“过严/与体验无关”，适度下调该项阈值或弱化权重；  
        - 对显著交互项（如 `solution × response_speed`）要求**联合通过**或分配更高合成权重；
        """)
        register_text("strategy_notes", "抽检策略与权重/联检建议已输出。")
    except Exception as e:
        st.warning(f"策略生成失败：{e}")

    # ============ 时间趋势 ============
    st.subheader("时间趋势分析（按月）")
    time_col = "质检时间" if "质检时间" in df.columns else None
    if time_col:
        dt_ser = pd.to_datetime(df[time_col], errors="coerce")
        df["month"] = dt_ser.dt.to_period("M").astype(str)
        trend = (df.dropna(subset=["month"])
                   .groupby("month")[["satisfied","overall_pass"]]
                   .mean().reset_index().sort_values("month"))
        trend["Satisfaction (%)"] = (trend["satisfied"]*100).round(2)
        trend["Pass Rate (%)"] = (trend["overall_pass"]*100).round(2)
        fig_trend, ax = plt.subplots(figsize=(9,4.5), dpi=140)
        x = np.arange(len(trend))
        ax.plot(x, trend["Satisfaction (%)"], marker="o", label="Satisfaction (%)")
        ax.plot(x, trend["Pass Rate (%)"], marker="o", label="Pass Rate (%)")
        ax.set_xticks(x); ax.set_xticklabels(trend["month"], rotation=30, ha="right")
        ax.set_title("Monthly Trend: Satisfaction vs Pass Rate"); ax.set_ylabel("Percentage (%)")
        ax.legend(); ax.grid(alpha=0.25, linestyle="--", linewidth=0.5)
        st.pyplot(fig_trend)
        register_fig("trend_lines", fig_trend)
        register_table("trend_table", trend)

    # ============ 分业务线 / 渠道 ============
    if "business_line" in df.columns:
        st.subheader("分业务线：整体通过率 vs 满意率")
        biz = (df.groupby("business_line")[pass_cols+["overall_pass","satisfied"]]
                 .mean().apply(lambda x: round(x*100,2)).reset_index())
        st.dataframe(biz[["business_line","overall_pass","satisfied"]])
        register_table("biz_overview", biz)
        fig_biz, ax_biz = plt.subplots(figsize=(8,4.5), dpi=140)
        sns.scatterplot(data=biz, x="overall_pass", y="satisfied", hue="business_line", s=120, ax=ax_biz)
        ax_biz.set_xlabel("Overall Pass Rate (%)"); ax_biz.set_ylabel("Satisfaction Rate (%)")
        ax_biz.set_title("Business Line: Pass vs Satisfaction")
        st.pyplot(fig_biz); register_fig("biz_scatter", fig_biz)

    if "ticket_channel" in df.columns:
        st.subheader("分渠道：整体通过率 vs 满意率")
        ch = (df.groupby("ticket_channel")[pass_cols+["overall_pass","satisfied"]]
                .mean().apply(lambda x: round(x*100,2)).reset_index())
        st.dataframe(ch[["ticket_channel","overall_pass","satisfied"]])
        register_table("channel_overview", ch)
        fig_ch, ax_ch = plt.subplots(figsize=(8,4.5), dpi=140)
        sns.scatterplot(data=ch, x="overall_pass", y="satisfied", hue="ticket_channel", s=120, ax=ax_ch)
        ax_ch.set_xlabel("Overall Pass Rate (%)"); ax_ch.set_ylabel("Satisfaction Rate (%)")
        ax_ch.set_title("Channel: Pass vs Satisfaction")
        st.pyplot(fig_ch); register_fig("channel_scatter", fig_ch)

    # ============ 自动结论 ============
    st.success("✅ 全部分析完成。")
    st.subheader("📊 自动结论（面向标准优化）")
    try:
        sig_items = coef_df[coef_df["p值"]<0.05]
        if not sig_items.empty:
            key_item = sig_items.sort_values("回归系数", ascending=False).iloc[0]["指标项"]
            lowest_item = sig_items.sort_values("回归系数", ascending=True).iloc[0]["指标项"]
            concl = f"""
**1️⃣ 最关键提升项：** `{key_item}` → 建议提高权重或细化二级准则（分档/示例库）。  
**2️⃣ 可能过严/定义模糊：** `{lowest_item}` → 建议降低阈值或从“必过”改为“加分项”。  
**3️⃣ 若出现“通过↑但满意↓”** → 标准偏离客户感知，优先复核 FP 区域样本并修订规则。  
**4️⃣ 对显著交互项** → 采用“联合通过”或“联动加权”，避免单点达标掩盖问题。  
"""
            st.markdown(concl)
            register_text("final_conclusion", concl)
        else:
            st.info("暂无显著项，建议扩大样本或延长观测窗口。")
    except Exception as e:
        st.warning(f"结论生成失败：{e}")

    # ====================== ⬇️ 导出报告与原始素材（DOCX / ZIP） ======================
    st.markdown("---")
    st.header("导出报告")
    st.caption("说明：若环境已安装 `python-docx` 将生成 Word 报告；同时提供 ZIP（包含 CSV 表格与 PNG 图片）。")

    colA, colB = st.columns(2)

    with colA:
        # 导出 ZIP（表格CSV + 图像PNG）
        if st.button("打包导出 ZIP（表+图）"):
            zip_buffer = BytesIO()
            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
                # CSVs
                for name, df_tbl in table_registry.items():
                    csv_bytes = df_tbl.to_csv(index=True).encode("utf-8-sig")
                    zf.writestr(f"tables/{name}.csv", csv_bytes)
                # PNG figures
                for name, fig in fig_registry.items():
                    img_bytes = BytesIO()
                    fig.savefig(img_bytes, format="png", dpi=200, bbox_inches="tight")
                    img_bytes.seek(0)
                    zf.writestr(f"figures/{name}.png", img_bytes.read())
                # 结论文本
                all_txt = "\n\n".join([f"[{k}]\n{text_registry[k]}" for k in text_registry])
                zf.writestr("summary/conclusions.txt", all_txt.encode("utf-8"))

            zip_buffer.seek(0)
            ts = dt.datetime.now().strftime("%Y%m%d_%H%M%S")
            st.download_button(
                label="下载 ZIP",
                data=zip_buffer,
                file_name=f"QC_Satisfaction_Report_{ts}.zip",
                mime="application/zip"
            )

    with colB:
        # 导出 DOCX（需要 python-docx）
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            if st.button("导出 Word（.docx）报告"):
                doc = Document()
                doc.add_heading("质检与满意度分析报告", 0)
                p = doc.add_paragraph(f"报告生成时间：{dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # 章节：结论
                doc.add_heading("一、自动结论（面向标准优化）", level=1)
                doc.add_paragraph(text_registry.get("final_conclusion", "（本次未生成显著性结论，建议扩大样本）"))

                # 章节：关键表格（精选Top）
                doc.add_heading("二、关键表格", level=1)
                key_tables = [
                    ("summary_rates", "各项通过率与满意率（%）"),
                    ("logit_single", "Logistic回归（单项）"),
                    ("interaction_ttest", "两两交互显著性"),
                    ("logit_interactions", "交互项回归（含系数）"),
                    ("weight_candidate", "权重候选（按回归系数归一）"),
                    ("trend_table", "月度趋势"),
                ]
                for key, title in key_tables:
                    if key in table_registry:
                        doc.add_heading(title, level=2)
                        df_tbl = table_registry[key]
                        # 写入表格
                        t = doc.add_table(rows=1, cols=len(df_tbl.columns)+1)
                        t.style = "Light List Accent 1"
                        hdr_cells = t.rows[0].cells
                        hdr_cells[0].text = ""
                        for i, c in enumerate(df_tbl.columns):
                            hdr_cells[i+1].text = str(c)
                        for idx, row in df_tbl.iterrows():
                            cells = t.add_row().cells
                            cells[0].text = str(idx)
                            for j, c in enumerate(df_tbl.columns):
                                cells[j+1].text = str(row[c])
                        doc.add_paragraph("")  # spacing

                # 章节：关键图像
                doc.add_heading("三、核心图表", level=1)
                key_figs = [
                    ("corr_matrix", "相关性热力图"),
                    ("logit_coef_bar", "Logistic 系数影响条形图"),
                    ("trend_lines", "月度趋势"),
                    ("biz_scatter", "业务线：通过率 vs 满意率"),
                    ("channel_scatter", "渠道：通过率 vs 满意率"),
                ]
                for key, title in key_figs:
                    if key in fig_registry:
                        doc.add_heading(title, level=2)
                        fig = fig_registry[key]
                        img_bytes = BytesIO()
                        fig.savefig(img_bytes, format="png", dpi=200, bbox_inches="tight")
                        img_bytes.seek(0)
                        doc.add_picture(img_bytes, width=Inches(6.2))
                        doc.add_paragraph("")

                # 保存到内存并下载
                doc_bytes = BytesIO()
                doc.save(doc_bytes)
                doc_bytes.seek(0)
                ts = dt.datetime.now().strftime("%Y%m%d_%H%M%S")
                st.download_button(
                    label="下载 Word 报告",
                    data=doc_bytes,
                    file_name=f"QC_Satisfaction_Report_{ts}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        except Exception as e:
            st.info(f"未检测到 python-docx 或导出失败：{e}\n如需 Word 导出：pip install python-docx")

else:
    st.info("请上传多个质检文件后开始分析。")
